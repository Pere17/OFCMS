import re
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

import os
BASE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(BASE, "chapters_3_4_5.md")
OUT = os.path.join(BASE, "OFCMS_Chapters_3_4_5.docx")

FONT = "Times New Roman"
BODY_SIZE = Pt(12)

def set_run_font(run, size=BODY_SIZE, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)

def add_body_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(10)
    # handle **bold** inline segments
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = p.add_run(part)
            set_run_font(run)
    return p

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(18 if level == 1 else 12)
    pf.space_after = Pt(12 if level == 1 else 8)
    pf.line_spacing = 1.5
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=Pt(15), bold=True)
        pf.page_break_before = True
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=Pt(13), bold=True)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=Pt(12), bold=True, italic=True)
    return p

def add_list_item(doc, text, ordered, level=0):
    style = 'List Number' if ordered else 'List Bullet'
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.left_indent = Inches(0.25 + 0.25 * level)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = p.add_run(part)
            set_run_font(run)
    return p

def add_code_block(doc, lines):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.3)
    for i, line in enumerate(lines):
        run = p.add_run(line + ("\n" if i < len(lines) - 1 else ""))
        run.font.name = "Consolas"
        run.font.size = Pt(10)

def set_cell_text(cell, text, bold=False, size=Pt(11)):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)

def add_table(doc, rows):
    # rows: list of list of strings; first row is header
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell_text_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
            set_cell_text(table.cell(r, c), cell_text_clean, bold=(r == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def parse_and_build():
    with open(SRC, encoding='utf-8') as f:
        raw_lines = f.read().split('\n')

    doc = Document()

    # Page setup: 1 inch margins, Letter/A4 default is fine
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # base style
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = BODY_SIZE

    i = 0
    n = len(raw_lines)
    first_heading_done = False

    while i < n:
        line = raw_lines[i]
        nxt = raw_lines[i + 1] if i + 1 < n else ""

        # blank line
        if line.strip() == "":
            i += 1
            continue

        # setext H1 (=== underline)
        if re.match(r'^=+$', nxt.strip()) and len(nxt.strip()) >= 3:
            add_heading(doc, line.strip(), level=1)
            i += 2
            continue

        # setext H2 (--- underline), but distinguish from markdown table separator / hr
        if re.match(r'^-{3,}$', nxt.strip()):
            add_heading(doc, line.strip(), level=2)
            i += 2
            continue

        # H3 via ###
        if line.startswith('### '):
            add_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue

        # code fence
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < n and not raw_lines[i].strip().startswith('```'):
                code_lines.append(raw_lines[i])
                i += 1
            i += 1  # skip closing fence
            add_code_block(doc, code_lines)
            continue

        # table block (starts with |)
        if line.strip().startswith('|'):
            table_lines = []
            while i < n and raw_lines[i].strip().startswith('|'):
                table_lines.append(raw_lines[i].strip())
                i += 1
            rows = []
            for idx, tl in enumerate(table_lines):
                if idx == 1 and re.match(r'^\|[\s\-:|]+\|$', tl):
                    continue  # separator row
                cells = [c.strip() for c in tl.strip('|').split('|')]
                rows.append(cells)
            if rows:
                add_table(doc, rows)
            continue

        # numbered list item
        m = re.match(r'^(\d+)\.\s+(.*)$', line.strip())
        if m:
            add_list_item(doc, m.group(2), ordered=True)
            i += 1
            continue

        # bullet list item
        m = re.match(r'^-\s+(.*)$', line.strip())
        if m:
            add_list_item(doc, m.group(1), ordered=False)
            i += 1
            continue

        # italic-only line (e.g., figure reference) treat as centered italic caption
        if line.strip().startswith('*(') and line.strip().endswith(')*'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip()[1:-1])
            set_run_font(run, italic=True, size=Pt(11))
            p.paragraph_format.space_after = Pt(10)
            i += 1
            continue

        # regular paragraph -- gather until blank line
        para_lines = [line]
        i += 1
        while i < n and raw_lines[i].strip() != "" and not re.match(r'^=+$', raw_lines[i].strip()) \
                and not (i + 1 < n and re.match(r'^-{3,}$', raw_lines[i + 1].strip())) \
                and not raw_lines[i].startswith('### ') \
                and not raw_lines[i].strip().startswith('|') \
                and not re.match(r'^\d+\.\s+', raw_lines[i].strip()) \
                and not raw_lines[i].strip().startswith('- ') \
                and not raw_lines[i].strip().startswith('```'):
            para_lines.append(raw_lines[i])
            i += 1
        add_body_paragraph(doc, " ".join(para_lines).strip())

    doc.save(OUT)
    print(f"Saved {OUT}")


if __name__ == "__main__":
    parse_and_build()
